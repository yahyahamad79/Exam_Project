# main.py
#yahya hamad
import streamlit as st
import fitz  # PyMuPDF للمعالجة الحقيقية
from docx import Document
import re
import io
import tempfile
import os


# 1. إعدادات الصفحة (يجب أن يكون هذا أول أمر برمي)
st.set_page_config(
    page_title="مولد الامتحانات الذكي",
    page_icon="📝",
    layout="wide"
)

# 2. تنسيق الواجهة (CSS) لدعم اللغة العربية
st.markdown("""
    <style>
        .stApp { direction: rtl; text-align: right; }
        .main-header {
            background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
            padding: 20px;
            border-radius: 10px;
            color: white;
            text-align: center;
            margin-bottom: 25px;
        }
        .question-card {
            background: #ffffff;
            padding: 15px;
            border-radius: 8px;
            margin: 10px 0;
            border-right: 5px solid #3b82f6;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
    </style>
""", unsafe_allow_html=True)

# 3. دالة استخراج الأسئلة (الخوارزمية المحدثة)
def extract_questions_from_pdf(text):
    questions = []
    # تقسيم النص إلى جمل بناءً على النقاط، الأسطر، والفاصلة المنقوطة
    sentences = re.split(r'[.\n؛]', text)
    count = 1
    
    # كلمات دالة على محتوى امتحاني
    key_words = ["يعرف", "تعد", "تعتبر", "يتكون", "هدف", "أهمية", "ميزة", "خصائص", "انواع", "عناصر"]
    
    for sentence in sentences:
        clean_text = sentence.strip()
        
        # قبول الجمل التي تزيد عن 20 حرفاً لضمان التقاط السطور العربية القصيرة
        if len(clean_text) > 20:
            is_educational = any(word in clean_text for word in key_words)
            
            if is_educational or len(clean_text) > 55:
                # تحديد المستوى تلقائياً
                level = "بسيط"
                if any(word in clean_text for word in ["بسبب", "بينما", "نتيجة"]):
                    level = "متوسط"
                if any(word in clean_text for word in ["التحوط", "الاستراتيجي", "تحليل"]):
                    level = "متقدم"
                
                # إضافة علامة استفهام إذا لم تكن موجودة
                q_text = clean_text + "؟" if not clean_text.endswith('؟') else clean_text
                
                questions.append({
                    "id": count,
                    "level": level,
                    "text": q_text
                })
                count += 1
                if count > 40: break
    return questions

# 4. دالة إنشاء ملف Word
def create_word_file(selected_questions):
    doc = Document()
    doc.add_heading('اختبار مادة الإدارة - مستخرج آلياً', 0)
    for i, q in enumerate(selected_questions, 1):
        p = doc.add_paragraph()
        p.add_run(f"السؤال {i}: ").bold = True
        p.add_run(q)
        doc.add_paragraph("الإجابة: ......................................................................")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 5. بناء واجهة المستخدم
def main():
    st.markdown('<div class="main-header"><h1>📝 مولد الامتحانات الذكي</h1><p>استخراج الأسئلة من ملفات PDF التعليمية</p></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.subheader("📤 الخطوة 1: ارفع الملف")
        file = st.file_uploader("اختر ملف PDF", type=['pdf'])
        if file:
            if st.button("🚀 استخراج الأسئلة", use_container_width=True):
                with st.spinner("جاري التحليل..."):
                    pdf_doc = fitz.open(stream=file.read(), filetype="pdf")
                    text = "".join([page.get_text() for page in pdf_doc])
                    
                    if text.strip():
                        st.session_state.qs = extract_questions_from_pdf(text)
                        st.success(f"تم استخراج {len(st.session_state.qs)} سؤالاً")
                    else:
                        st.error("لم نتمكن من قراءة نص من الملف. تأكد أنه ليس صوراً.")

    with col2:
        st.subheader("📋 الخطوة 2: مراجعة الأسئلة")
        if 'qs' in st.session_state and st.session_state.qs:
            final_list = []
            for q in st.session_state.qs:
                if st.checkbox(f"[{q['level']}] {q['text']}", key=f"q_{q['id']}", value=True):
                    final_list.append(q['text'])
            
            st.divider()
            if st.button("📄 تحميل كملف Word", type="primary"):
                if final_list:
                    word_data = create_word_file(final_list)
                    st.download_button("اضغط للتحميل الآن", word_data, "My_Exam.docx")
                else:
                    st.warning("اختر سؤالاً واحداً على الأقل")
        else:
            st.info("النتائج ستظهر هنا بعد رفع الملف والضغط على استخراج.")

if __name__ == "__main__":
    main()
